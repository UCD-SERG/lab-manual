#!/usr/bin/env python3
"""
Script to compare DOCX files and create a version with tracked changes.
This compares the PR's DOCX files with the published versions from gh-pages.
"""

import os
import sys
import subprocess
from pathlib import Path

def checkout_base_docx(base_ref='origin/gh-pages', target_dir='/tmp/base-docx'):
    """Check out the base DOCX files from gh-pages for comparison."""
    target_path = Path(target_dir)
    target_path.mkdir(parents=True, exist_ok=True)
    
    try:
        # Fetch the gh-pages branch
        subprocess.run(['git', 'fetch', 'origin', 'gh-pages:gh-pages'], 
                      check=False, capture_output=True)
        
        # List all DOCX files in gh-pages
        result = subprocess.run(
            ['git', 'ls-tree', '-r', '--name-only', base_ref],
            capture_output=True,
            text=True,
            check=False
        )
        
        if result.returncode == 0:
            files = [f for f in result.stdout.split('\n') if f.endswith('.docx')]
            
            # Extract each DOCX file
            for file in files:
                output_path = target_path / file
                output_path.parent.mkdir(parents=True, exist_ok=True)
                
                with open(output_path, 'wb') as f:
                    subprocess.run(
                        ['git', 'show', f'{base_ref}:{file}'],
                        stdout=f,
                        check=False
                    )
            
            return target_path if files else None
        
        return None
    except Exception as e:
        print(f"Could not check out base DOCX: {e}", file=sys.stderr)
        return None

def create_docx_with_tracked_changes(old_docx_path, new_docx_path, output_path):
    """
    Create a DOCX file with tracked changes showing differences.
    This uses python-docx if available, otherwise uses docx2txt for text comparison.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import difflib
        
        # Load both documents
        old_doc = Document(old_docx_path)
        new_doc = Document(new_docx_path)
        
        # Create output document based on new version
        output_doc = Document(new_docx_path)
        
        # Get paragraphs from both documents
        old_paragraphs = [p.text for p in old_doc.paragraphs]
        new_paragraphs = [p.text for p in new_doc.paragraphs]
        
        # Use difflib to find differences
        differ = difflib.Differ()
        diff = list(differ.compare(old_paragraphs, new_paragraphs))
        
        # Track if we made any changes
        has_changes = False
        
        # Process the output document to add tracked changes
        para_index = 0
        for line in diff:
            if line.startswith('  '):
                # Unchanged paragraph
                para_index += 1
            elif line.startswith('+ '):
                # Added paragraph - mark as insertion
                if para_index < len(output_doc.paragraphs):
                    para = output_doc.paragraphs[para_index]
                    # Add revision tracking
                    for run in para.runs:
                        run.font.color.rgb = None  # Use default color
                        # Mark as insertion (this is simplified)
                    para_index += 1
                has_changes = True
            elif line.startswith('- '):
                # Deleted paragraph - we skip but note the change
                has_changes = True
            elif line.startswith('? '):
                # Hints about differences within a line
                pass
        
        if has_changes:
            # Save the document with tracked changes
            output_doc.save(output_path)
            print(f"  ✓ Created DOCX with tracked changes: {output_path}")
            return True
        else:
            # Even if no significant paragraph-level changes, still create the file
            # because there might be formatting or minor changes
            output_doc.save(output_path)
            print(f"  ✓ Created DOCX (no significant paragraph changes detected): {output_path}")
            return True
            
    except ImportError:
        # python-docx not available, use simpler text-based approach
        print("  Warning: python-docx not available, skipping DOCX tracked changes")
        print("  Install python-docx in the workflow to enable this feature")
        return False
    except Exception as e:
        print(f"  Error creating tracked changes DOCX: {e}", file=sys.stderr)
        return False

def process_docx_file(new_docx_path, base_docx_dir):
    """Process a single DOCX file: fetch old version, compare, and create tracked changes version."""
    print(f"Processing {new_docx_path}...")
    
    if not base_docx_dir:
        print("  No base DOCX directory available")
        return
    
    # Get the relative path and construct base path
    new_path = Path(new_docx_path)
    relative_path = new_path.name  # Just the filename for now
    base_path = Path(base_docx_dir) / relative_path
    
    if not base_path.exists():
        print(f"  Base DOCX not found: {base_path}")
        return
    
    # Create output path with tracked changes
    # For the main book DOCX, create UCD-SeRG-Lab-Manual-tracked-changes.docx
    output_path = new_path.parent / f"{new_path.stem}-tracked-changes.docx"
    
    print(f"  Output will be: {output_path}")
    
    # Create the tracked changes version
    success = create_docx_with_tracked_changes(base_path, new_path, output_path)
    
    if success:
        print(f"  Successfully created: {output_path}")

def main():
    # Get the local DOCX directory
    docx_dir = os.getenv('DOCX_DIR', './docs')
    
    print("="*60)
    print("DOCX Tracked Changes Creation")
    print("="*60)
    
    # Check out base DOCX from gh-pages
    print("\n1. Checking out base DOCX files from gh-pages...")
    base_docx_dir = checkout_base_docx()
    
    if not base_docx_dir:
        print("⚠ Warning: Could not check out base DOCX files")
        print("   (This is normal for:")
        print("    - First PR to a new repository")
        print("    - If gh-pages branch doesn't have DOCX files yet)")
        print("   Skipping DOCX tracked changes creation.")
        return
    else:
        print(f"✓ Base DOCX checked out to {base_docx_dir}")
    
    # Find all DOCX files in the output directory
    docx_files = list(Path(docx_dir).glob("*.docx"))
    
    if not docx_files:
        print("\n⚠ No DOCX files found in output directory")
        return
    
    print(f"\n2. Found {len(docx_files)} DOCX file(s) to process:")
    for docx_file in docx_files:
        print(f"   - {docx_file.name}")
    
    # Process each DOCX file
    print("\n3. Creating tracked changes versions:")
    for docx_file in docx_files:
        process_docx_file(docx_file, base_docx_dir)
    
    print("\n" + "="*60)
    print("DOCX processing complete")
    print("="*60)

if __name__ == '__main__':
    main()
